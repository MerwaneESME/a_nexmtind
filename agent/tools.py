"""Outils métier pour l'agent BTP V2 - Avec function calling."""
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Any, List, Literal, Optional

import docx
import pypdf
from langchain_core.tools import tool
from pydantic import BaseModel, Field

from .config import InvoiceSchema, QuoteSchema
from .supabase_client import get_client

OUTPUT_DIR = Path(__file__).parent.parent / "output"


def _normalize_decimal(value: Any, default: Decimal = Decimal(0)) -> Decimal:
    try:
        return Decimal(str(value))
    except (InvalidOperation, TypeError):
        return default


# ==================== Extraction de fichiers ====================

class ExtractInput(BaseModel):
    file_path: str
    doc_type: Literal["quote", "invoice", "auto"] = "auto"


@tool(args_schema=ExtractInput)
def extract_pdf_tool(file_path: str, doc_type: str = "auto"):
    """Extraction texte/structure depuis PDF/DOCX/PNG/JPG."""
    path = Path(file_path)
    if not path.exists():
        return {"error": "file_not_found", "file_path": file_path}

    suffix = path.suffix.lower()
    text = ""
    page_count = 0
    
    if suffix == ".pdf":
        reader = pypdf.PdfReader(path)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        page_count = len(reader.pages)
    elif suffix == ".docx":
        document = docx.Document(path)
        text = "\n".join(p.text for p in document.paragraphs)
        page_count = len(document.paragraphs)
    elif suffix in (".png", ".jpg", ".jpeg"):
        try:
            import pytesseract
            from PIL import Image
            image = Image.open(path)
            text = pytesseract.image_to_string(image)
            page_count = 1
        except ImportError:
            return {"error": "ocr_unavailable", "file_path": file_path}
    else:
        return {"error": "unsupported_format", "file_path": file_path}

    detected = doc_type if doc_type != "auto" else "quote"
    return {
        "doc_type": detected,
        "parsed_text": text,
        "page_count": page_count,
        "file_path": str(path),
    }


# ==================== Nettoyage et calculs ====================

class CleanLinesInput(BaseModel):
    lines: List[dict] = Field(default_factory=list)
    default_vat_rate: float | None = 20.0


@tool(args_schema=CleanLinesInput)
def clean_lines_tool(lines: List[dict], default_vat_rate: float | None = 20.0):
    """Nettoie les lignes (quantités, unités, prix) et déduplication."""
    cleaned = []
    warnings = []
    seen_desc = set()

    for idx, line in enumerate(lines):
        desc = str(line.get("description") or "").strip()
        if not desc:
            warnings.append({"index": idx, "issue": "description_vide"})
            continue
        
        key = desc.lower()
        if key in seen_desc:
            warnings.append({"index": idx, "issue": "duplicate_description", "description": desc})
        seen_desc.add(key)

        qty = _normalize_decimal(line.get("quantity") or line.get("qty") or 0)
        price = _normalize_decimal(line.get("unit_price_ht") or line.get("unit_price") or 0)
        vat = _normalize_decimal(line.get("vat_rate", default_vat_rate if default_vat_rate is not None else 20))
        discount = _normalize_decimal(line.get("discount_rate") or 0)
        unit = (line.get("unit") or "").strip() or None

        if qty < 0:
            warnings.append({"index": idx, "issue": "negative_quantity"})
            qty = abs(qty)
        if price < 0:
            warnings.append({"index": idx, "issue": "negative_price"})
            price = abs(price)
        if vat < 0:
            warnings.append({"index": idx, "issue": "negative_vat_rate"})
            vat = abs(vat)

        cleaned.append({
            "description": desc,
            "quantity": float(qty),
            "unit": unit,
            "unit_price_ht": float(price),
            "vat_rate": float(vat),
            "discount_rate": float(discount),
        })

    return {"lines": cleaned, "warnings": warnings}


class CalculateTotalsInput(BaseModel):
    lines: List[dict] = Field(default_factory=list)
    default_vat_rate: float | None = 20.0
    doc_type: Literal["quote", "invoice"] | None = None


@tool(args_schema=CalculateTotalsInput)
def calculate_totals_tool(lines: List[dict], default_vat_rate: float | None = 20.0, doc_type: str | None = None):
    """Calcule HT/TVA/TTC et signale les incohérences numériques."""
    total_ht = Decimal("0")
    total_tva = Decimal("0")
    issues = []

    for idx, line in enumerate(lines):
        qty = _normalize_decimal(line.get("quantity") or line.get("qty") or 0)
        price = _normalize_decimal(line.get("unit_price_ht") or line.get("unit_price") or 0)
        vat_rate = _normalize_decimal(line.get("vat_rate", default_vat_rate if default_vat_rate is not None else 20))
        discount = _normalize_decimal(line.get("discount_rate") or 0)

        line_total_ht = qty * price * (Decimal(1) - discount / Decimal(100))
        line_total_tva = line_total_ht * vat_rate / Decimal(100)

        if qty == 0 or price == 0:
            issues.append({"index": idx, "issue": "zero_value_line", "severity": "medium"})
        if vat_rate == 0:
            issues.append({"index": idx, "issue": "vat_missing_or_zero", "severity": "low"})

        total_ht += line_total_ht
        total_tva += line_total_tva

    totals = {
        "total_ht": float(total_ht),
        "total_tva": float(total_tva),
        "total_ttc": float(total_ht + total_tva),
    }
    return {"totals": totals, "issues": issues, "doc_type": doc_type or "quote"}


# ==================== Validation ====================

class ValidateInput(BaseModel):
    payload: dict


@tool(args_schema=ValidateInput)
def validate_devis_tool(payload: dict):
    """Contrôle TVA, mentions obligatoires et cohérence des totaux."""
    doc_type = payload.get("doc_type") or "quote"
    try:
        document = InvoiceSchema(**payload) if doc_type == "invoice" else QuoteSchema(**payload)
    except Exception as exc:
        return {"valid": False, "errors": [str(exc)], "issues": [], "totals": {}}

    totals = document.totals()
    issues = []

    # Vérifier cohérence totaux
    for key, computed_value in totals.items():
        declared = payload.get(key)
        if declared is not None:
            diff = abs(_normalize_decimal(declared) - computed_value)
            if diff > Decimal("0.01"):
                issues.append({
                    "field": key,
                    "issue": f"ecart_total_{key}",
                    "details": {"declared": float(_normalize_decimal(declared)), "computed": float(computed_value)},
                    "severity": "medium",
                })

    # Mentions obligatoires
    if not document.payment_terms:
        issues.append({"field": "payment_terms", "issue": "conditions_paiement_manquantes", "severity": "high"})
    
    if doc_type == "invoice":
        if not document.penalties_late_payment:
            issues.append({"field": "penalties_late_payment", "issue": "penalites_retard_manquantes", "severity": "high"})
        if not document.professional_liability:
            issues.append({"field": "professional_liability", "issue": "mention_rc_pro_manquante", "severity": "medium"})
        if not document.due_date:
            issues.append({"field": "due_date", "issue": "date_echeance_manquante", "severity": "high"})

    # Vérifier lignes
    for idx, line in enumerate(document.line_items):
        if line.vat_rate < 0:
            issues.append({"index": idx, "field": "vat_rate", "issue": "tva_negative", "severity": "medium"})
        if line.vat_rate == 0:
            issues.append({"index": idx, "field": "vat_rate", "issue": "tva_absente", "severity": "low"})
        if line.quantity == 0 or line.unit_price_ht == 0:
            issues.append({"index": idx, "field": "line_items", "issue": "ligne_zero", "severity": "medium"})

    return {
        "valid": len([i for i in issues if i.get("severity") == "high"]) == 0,
        "issues": issues,
        "totals": {k: float(v) for k, v in totals.items()},
    }


# ==================== Recherche Supabase ====================

class SupabaseLookupInput(BaseModel):
    query: Optional[str] = None
    mode: Literal["clients", "materials", "history", "prefill", "auto"] = "auto"
    limit: int = 10


@tool(args_schema=SupabaseLookupInput)
def supabase_lookup_tool(query: Optional[str] = None, mode: str = "auto", limit: int = 10):
    """Recherche clients/matériaux/historiques dans Supabase."""
    sb = get_client()
    if not sb:
        return {"results": {}, "error": "supabase_not_configured"}

    results: dict[str, Any] = {}
    try:
        if mode in ("clients", "auto", "prefill"):
            q = sb.table("clients").select("id,name,address,contact").order("created_at", desc=True).limit(limit)
            if query:
                q = q.ilike("name", f"%{query}%")
            results["clients"] = q.execute().data or []

        if mode in ("materials", "auto", "prefill"):
            q_mat = sb.table("devis_items").select("description,unit_price,qty").limit(limit)
            mats = q_mat.execute().data or []
            dedup = []
            seen = set()
            for item in mats:
                desc = item.get("description")
                if desc and desc not in seen:
                    seen.add(desc)
                    dedup.append({"description": desc, "unit_price": item.get("unit_price"), "qty": item.get("qty")})
            results["materials"] = dedup

        if mode in ("history", "auto", "prefill"):
            q_hist = sb.table("devis").select("id,client_id,total,status,metadata").order("created_at", desc=True).limit(limit)
            results["history"] = q_hist.execute().data or []
    except Exception as exc:
        return {"results": results, "error": str(exc)}

    return {"results": results}


# ==================== Liste des outils disponibles ====================

AVAILABLE_TOOLS = [
    extract_pdf_tool,
    clean_lines_tool,
    calculate_totals_tool,
    validate_devis_tool,
    supabase_lookup_tool,
]
